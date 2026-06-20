#!/usr/bin/env python3
"""
render_report.py — render a CSB-style markdown report to DOCX (and optionally PDF).

Self-contained and defensive: it uses whatever is installed and tells you exactly
what it produced. No tool is mandatory; DOCX alone satisfies the requested output.

Conversion strategy
  DOCX:  pandoc  ->  python-docx (fallback)
  PDF:   pandoc (needs a LaTeX engine)  ->  libreoffice --headless (DOCX -> PDF)

USAGE
  python3 render_report.py report.md                # -> report.docx
  python3 render_report.py report.md --out out_dir  # write into out_dir/
  python3 render_report.py report.md --pdf          # also try to produce a PDF
"""

import argparse
import os
import shutil
import subprocess
import sys


def have(tool: str) -> bool:
    return shutil.which(tool) is not None


def run(cmd: list, timeout: int = 120) -> bool:
    """Run a command; return True only if it exits 0 within the timeout."""
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL,
                       stderr=subprocess.DEVNULL, timeout=timeout)
        return True
    except (subprocess.CalledProcessError, OSError, subprocess.TimeoutExpired):
        return False


def wrote(path: str) -> bool:
    """A converter only 'succeeded' if it actually produced a non-empty file.

    Some tools (notably libreoffice --headless) exit 0 without writing output."""
    return os.path.exists(path) and os.path.getsize(path) > 0


def docx_via_pandoc(md: str, out_docx: str) -> bool:
    if not have("pandoc"):
        return False
    return run(["pandoc", md, "-o", out_docx]) and wrote(out_docx)


def docx_via_python_docx(md: str, out_docx: str) -> bool:
    """Minimal markdown -> docx: headings, tables, bullets, paragraphs."""
    try:
        from docx import Document
    except ImportError:
        return False

    doc = Document()
    with open(md, encoding="utf-8") as f:
        lines = f.read().splitlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            doc.add_heading(text, level=min(level, 9))
            i += 1
            continue

        # Tables (consecutive lines starting with "|")
        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            rows = []
            for r in block:
                cells = [c.strip() for c in r.strip("|").split("|")]
                # skip separator rows like |---|---|
                if all(set(c) <= set("-: ") and c for c in cells):
                    continue
                rows.append(cells)
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=0, cols=ncols)
                try:
                    table.style = "Light Grid Accent 1"
                except Exception:
                    pass
                for r in rows:
                    cells = table.add_row().cells
                    for j, val in enumerate(r):
                        cells[j].text = val
            continue

        # Bullets
        if stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            doc.add_paragraph(stripped.lstrip("> ").strip(), style="Intense Quote")
            i += 1
            continue

        doc.add_paragraph(stripped)
        i += 1

    try:
        doc.save(out_docx)
        return True
    except Exception:
        return False


def pdf_via_pandoc(md: str, out_pdf: str) -> bool:
    if not have("pandoc"):
        return False
    # pandoc needs a PDF engine (LaTeX, wkhtmltopdf, etc.)
    return run(["pandoc", md, "-o", out_pdf]) and wrote(out_pdf)


def pdf_via_libreoffice(docx: str, out_dir: str) -> bool:
    soffice = "libreoffice" if have("libreoffice") else ("soffice" if have("soffice") else None)
    if not soffice or not os.path.exists(docx):
        return False
    out_pdf = os.path.join(out_dir, os.path.splitext(os.path.basename(docx))[0] + ".pdf")
    # A private profile avoids clashing with any running soffice instance.
    profile = "-env:UserInstallation=file://" + os.path.join(out_dir, ".lo_profile")
    run([soffice, "--headless", "--norestore", profile,
         "--convert-to", "pdf", "--outdir", out_dir, docx])
    return wrote(out_pdf)  # libreoffice may exit 0 without writing; trust the file


def main() -> int:
    ap = argparse.ArgumentParser(description="Render a CSB markdown report to DOCX/PDF.")
    ap.add_argument("markdown", help="Path to the markdown report")
    ap.add_argument("--out", default=None, help="Output directory (default: alongside input)")
    ap.add_argument("--pdf", action="store_true", help="Also attempt a PDF")
    args = ap.parse_args()

    md = args.markdown
    if not os.path.isfile(md):
        print(f"ERROR: not found: {md}", file=sys.stderr)
        return 2

    out_dir = args.out or os.path.dirname(os.path.abspath(md))
    os.makedirs(out_dir, exist_ok=True)
    base = os.path.splitext(os.path.basename(md))[0]
    out_docx = os.path.join(out_dir, base + ".docx")
    out_pdf = os.path.join(out_dir, base + ".pdf")

    produced = []

    # DOCX
    if docx_via_pandoc(md, out_docx):
        produced.append(("DOCX", out_docx, "pandoc"))
    elif docx_via_python_docx(md, out_docx):
        produced.append(("DOCX", out_docx, "python-docx"))
    else:
        print("WARNING: could not produce DOCX. Install pandoc or `pip install python-docx`.",
              file=sys.stderr)

    # PDF (optional)
    if args.pdf:
        if pdf_via_pandoc(md, out_pdf):
            produced.append(("PDF", out_pdf, "pandoc"))
        elif os.path.exists(out_docx) and pdf_via_libreoffice(out_docx, out_dir):
            produced.append(("PDF", out_pdf, "libreoffice"))
        else:
            print("WARNING: could not produce PDF. Install pandoc+LaTeX or libreoffice. "
                  "DOCX (if produced) still satisfies the output requirement.", file=sys.stderr)

    if not produced:
        print("ERROR: no artifacts produced.", file=sys.stderr)
        return 1

    for kind, path, engine in produced:
        print(f"Produced {kind}: {path}  (via {engine})")
    return 0


if __name__ == "__main__":
    sys.exit(main())
